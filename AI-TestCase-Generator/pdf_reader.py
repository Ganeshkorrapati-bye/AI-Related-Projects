"""Document ingestion: extract plain text from PDF, DOCX, and TXT uploads.

Each reader is defensive: unsupported formats, corrupt files, and empty
documents raise :class:`DocumentReadError` with a user-friendly message.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

from config import MAX_REQUIREMENT_CHARS, MAX_UPLOAD_MB

logger = logging.getLogger(__name__)


class DocumentReadError(Exception):
    """Raised when an uploaded document cannot be read or is invalid."""


def _check_size(data: bytes) -> None:
    size_mb = len(data) / (1024 * 1024)
    if size_mb > MAX_UPLOAD_MB:
        raise DocumentReadError(
            f"File is {size_mb:.1f} MB which exceeds the {MAX_UPLOAD_MB} MB limit."
        )


def _read_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency guaranteed by requirements
        raise DocumentReadError("PDF support requires the 'pdfplumber' package.") from exc

    text_parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                if extracted:
                    text_parts.append(extracted)
    except Exception as exc:
        raise DocumentReadError("The PDF file appears to be corrupt or unreadable.") from exc
    return "\n".join(text_parts)


def _read_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentReadError("DOCX support requires the 'python-docx' package.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentReadError("The Word document appears to be corrupt or unreadable.") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError("Unable to decode the text file. Please use UTF-8 encoding.")


_READERS = {
    "pdf": _read_pdf,
    "docx": _read_docx,
    "txt": _read_txt,
}


def extract_text(data: bytes, filename: str) -> str:
    """Extract and normalise text from raw upload bytes.

    Args:
        data: Raw file bytes.
        filename: Original filename (used only for its extension).

    Returns:
        Extracted, whitespace-normalised text.

    Raises:
        DocumentReadError: For unsupported formats, oversized, empty, or corrupt
            documents.
    """
    if not data:
        raise DocumentReadError("The uploaded file is empty.")

    _check_size(data)

    extension = Path(filename).suffix.lower().lstrip(".")
    reader = _READERS.get(extension)
    if reader is None:
        supported = ", ".join(sorted(_READERS))
        raise DocumentReadError(
            f"Unsupported format '.{extension}'. Supported formats: {supported}."
        )

    logger.info("Extracting text from %s (%d bytes)", filename, len(data))
    text = reader(data).strip()

    if not text:
        raise DocumentReadError(
            "No readable text was found. Scanned/image-only PDFs are not supported."
        )

    if len(text) > MAX_REQUIREMENT_CHARS:
        logger.warning(
            "Document truncated from %d to %d characters", len(text), MAX_REQUIREMENT_CHARS
        )
        text = text[:MAX_REQUIREMENT_CHARS]

    return text
